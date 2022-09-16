# ОБРАБОТКА ТАБЛИЦЫ И ЗАПИСЬ В ТХТ-ФАЙЛ

from docx import Document
doc = Document('data2.docx')
table = doc.tables[0]

data_in_table = []

# проходимся по строкам таблицы `i`
for i, row in enumerate(table.rows):
    row_data = {}
    # проходимся по ячейкам таблицы `i` и строки `j`
    for j, cell in enumerate(row.cells):
        if i == 0: # шапка таблицы не нужна
            break
        # добавляем значение ячейки в соответствующий
        # список, созданного словаря под данные таблиц
        key_tuple = ('Num', 'Date', 'View', 'Description', 'Scale', 'Damage')
        if j != 0:
            row_data[key_tuple[j]] = cell.text
        else:
            row_data[key_tuple[0]] = f'{i}'

    if len(row_data) != 0:
        data_in_table.append(row_data)

with open('save_data2.txt','w', encoding="utf-8") as f:
    f.write(str(data_in_table))

# ___________________________________________________
# ЧТЕНИЕ ВСЕХ ТАБЛИЦ ВОРДА
# # полноценный рабочий код для всех таблиц документа
# from docx import Document
#
# doc = Document('data.docx')
# # последовательность всех таблиц документа
# all_tables = doc.tables
# print(all_tables)
# print('Всего таблиц в документе:', len(all_tables))
#
# # создаем пустой словарь под данные таблиц
# data_tables = {i:None for i in range(len(all_tables))}
# # проходимся по таблицам
# for i, table in enumerate(all_tables):
#     print('\nДанные таблицы №', i)
#     # создаем список строк для таблицы `i` (пока пустые)
#     data_tables[i] = [[] for _ in range(len(table.rows))]
#     # проходимся по строкам таблицы `i`
#     for j, row in enumerate(table.rows):
#         # проходимся по ячейкам таблицы `i` и строки `j`
#         for cell in row.cells:
#             # добавляем значение ячейки в соответствующий
#             # список, созданного словаря под данные таблиц
#             data_tables[i][j].append(cell.text)
#
#     # смотрим извлеченные данные
#     # (по строкам) для таблицы `i`
#     print(data_tables[i])
#     print('\n')
#
# print('Данные всех таблиц документа:')
# print(data_tables)

# ЧТЕНИЕ ФАЙЛА
# with open('save_data.txt', 'r', encoding="utf-8") as f:
#     data = f.read()
# print(data)
# print(type(data))
#
# print(eval(data))
# print(type(eval(data)))